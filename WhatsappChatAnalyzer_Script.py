#!/usr/bin/env python
# coding: utf-8

# In[1]:


#Import required libraries
import re
import os
import regex
import pandas as pd
import numpy as np
import emoji
import plotly.express as px
from collections import Counter
import nltk 
from nltk.corpus import stopwords
nltk.download('stopwords')
STOPWORDS = stopwords.words('english')
from datetime import datetime, time, date, timedelta
import matplotlib.pyplot as plt
from PIL import Image
from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator
get_ipython().run_line_magic('matplotlib', 'inline')
from whatstk import WhatsAppChat
from docx import Document
from docx.shared import Cm
from reportlab.pdfgen.canvas import Canvas


# In[2]:


#Define content of varibales months and day
months= ['January', 'February', 'March', 'April', 'May', 'June', 'July', 'August', 'September', 'October', 'November', 'December']
days = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday', 'Sunday']


# In[3]:


#Definde the function whatsapp_chat_processor - Import .txt whatsapp data and create a dataframe for further analysis

def whatsapp_chat_processor():
    global filepath
    filepath = str(input('Please Enter the File Path of Your Whatsapp Chat: ')) #
    #Convert the .txt file to a readable dataframe
    try:
        chat = WhatsAppChat.from_source(filepath=filepath)
        chat = chat.df
    except:
        print('File not Found Error - Please provide the correct path') #error message if non correct filepath
        filepath = str(input('Please Enter the File Path of Your Whatsapp Chat: '))
        chat = WhatsAppChat.from_source(filepath=filepath)
        chat = chat.df
       
    #create new columns in the date frame
    #comlumn date   
    chat['date'] = pd.to_datetime(chat['date'])
    #column weekday
    chat['weekday'] = chat['date'].apply(lambda x: x.day_name())
    #column month_sent
    chat['month_sent'] = chat['date'].apply(lambda x: x.month_name()) 
    #column date
    chat['dateTime'] = [d.date() for d in chat['date']] 
    #column hour
    chat['hour'] = [d.time().hour for d in chat['date']]
    #column urlcount
    URLPATTERN = r'(https?://\S+)'
    chat['urlcount'] = chat.message.apply(lambda x: re.findall(URLPATTERN, x)).str.len()
    #column Letter_Count
    chat['Letter_Count'] = chat['message'].apply(lambda s : len(s))
    #column Word_Count
    chat['Word_Count'] = chat['message'].apply(lambda s : len(s.split(' ')))
    #Emojis
    chat["emoji"] = chat["message"].apply(get_emojis)
    return chat #return a datafram, for the further analysis


# In[4]:


#If user displays the user list (inserted continue) and inerts specific names
def author_specific_statistics(folder, words_authors, author_name_list):
    
    author_name = str(input('Please provide a List of specific Usernames (seperated by a comma):'))
    author_name = author_name.split(', ') #create list with inputet usernames
    print(author_name)
    for i in author_name: #create for each name in the list a graph with most frequently used words
        if i.lower() in author_name_list:
            dummy_df = words_authors[words_authors['username'].str.lower() == i.lower()]
            author_fig = px.bar(dummy_df, x='words', y='count',
                     labels={'words':'Most Common Words', 'count':'Count'}, 
                     height=380)
            author_fig.update_traces(marker_color='#EDCC8B', marker_line_color='#D4A29C',
                          marker_line_width=1.5, opacity=0.6)
            author_fig.update_layout(title_text='{} Commond Words '.format(i))
            author_fig.write_image(folder + '\\most_common_words_author_{}_{}.png'
                                        .format(re.sub('[^A-Za-z0-9]+', '',str(i)),datetime.now()
                                                .strftime('%d_%m_%Y')), format='.png'
                                         )
        else:
            print('Username not in Database - Please Enter a correct Username') #if a wrong username is inserted
            author_specific_statistics(folder, words_authors, author_name_list)


# In[5]:


#Define the function most_common_words
def most_common_words(chat, folder):
    words = chat[['username','message']].copy()
    stopwords = list(STOPWORDS)
    extra = ["<multimedia", "omitido>", "k", "d","si","multimedia", "omitido"] #create additional list
    stopwords = stopwords + extra #combine to new stopword list
    words["message"] = (words["message"] #prepare the column message (convert to smaller, remove stop words)
                               .str.lower()
                               .str.split()
                               .apply(lambda x: [item for item in x if item not in stopwords])
                               .explode()
                               .reset_index(drop=True)
                     )
    words['message']= words['message'].replace('nan', np.NaN)
    words['message']= words['message'].replace('', np.NaN)

    words_dict = dict(Counter(words.message))
    words_dict = sorted(words_dict.items(), key=lambda x: x[1], reverse=True)

    words_dict = pd.DataFrame(words_dict, columns=['words', 'count']); #create a data frame with the frequency per word

    #create a plot with the most common words and save as picture in folder
    most_common_words_fig = px.bar(words_dict.head(10).dropna(), x='words', y='count',
                     labels={'words':'Most Common Words', 'count':'Count'}, 
                     height=400)
    most_common_words_fig.update_traces(marker_color='#EDCC8B', marker_line_color='#D4A29C',
                      marker_line_width=1.5, opacity=0.6)
    most_common_words_fig.update_layout(title_text='Most Commond Words')
    most_common_words_fig.write_image(folder + '\\most_common_words_{}.png'
                                    .format(datetime.now()
                                            .strftime('%d_%m_%Y')), format='.png'
                                     )
    
    #prepare data for further analysis (count words by username)
    words_authors =  (words.set_index('username')['message']
                      .dropna()
                      .groupby(level=0)
                      .value_counts()
                      .groupby(level=0)
                      .head()
                      .rename_axis(('username','words'))
                      .reset_index(name='count'))
    
    #Input by the user - For how many people he or she would like to display the statistics
    author_input = input('Please Enter an Integer to see the most common Words by Users or enter Continue to provide a specific list of Usernames \n (e.g. 5 will show you the most common words of the first 5 User): ')
    try:
        author_input = int(author_input) #test if number has been inserted
    except:
        author_input = author_input
        
    if isinstance(author_input, int) == True:
        print('Integer entered')
        l = words_authors.username.unique()[:int(author_input)] #create a list with all usernames
       
    #display for selected number of or specified user name the most frequent words
        for i in range(len(l)):
            dummy_df = words_authors[words_authors['username'] == l[i]]
            author_fig = px.bar(dummy_df, x='words', y='count',
                         labels={'words':'Most Common Words', 'count':'Count'}, 
                         height=380)
            author_fig.update_traces(marker_color='#EDCC8B', marker_line_color='#D4A29C',
                          marker_line_width=1.5, opacity=0.6)
            author_fig.update_layout(title_text='{} Commond Words '.format(l[i]))
            author_fig.write_image(folder + '\\most_common_words_author_{}_{}.png'
                                        .format(re.sub('[^A-Za-z0-9]+', '',str(l[i])),datetime.now()
                                                .strftime('%d_%m_%Y')), format='.png'
                                         )
    #Print a list with all usernames of the WhatsApp group, if user entered "continue"
    elif author_input.lower() == 'continue': 
        author_name_list = [i.encode('ascii','ignore').decode('UTF-8').lower() for i in set(chat.username)]
        print('Usernames found in the Chat: \n',author_name_list)
        #use funtion owner spectic statistics (form above), based on the provided list of the user
        author_specific_statistics(folder = folder, words_authors = words_authors, author_name_list = author_name_list)


# In[6]:


#function to extract all the emojis from a text (here message)
def get_emojis(text):
    emoji_list = []
    data = regex.findall(r'\X', text)
    for word in data:
        if any(char in emoji.UNICODE_EMOJI for char in word):
            emoji_list.append(word)
    return emoji_list


# In[7]:


#function to remove the emojis from the text
def emoji_free_text(text):
    return emoji.get_emoji_regexp().sub(r'', text)

#function to remove the urls from the text
def remove_urls(text):
    url_pattern = re.compile(r'https?://\S+|www\.\S+')
    return url_pattern.sub(r'', text)


# In[8]:


#Creat a report in Word
def create_report_docx(total_messages, media_messages, average_message_words, average_message_letters, average_message_day, filepath, folder):
    doc = Document('WhatsApp Chat Report.docx')
    for p in doc.paragraphs:
        p.text = p.text.replace('WhatsApp Chat File:', 'WhatsApp Chat File: {}'.format(filepath))
        p.text = p.text.replace('Total Messages:', 'Total Messages: {}'.format(total_messages))
        p.text = p.text.replace('Media Message:', 'Media Message: {}'.format(media_messages))
        p.text = p.text.replace('Average Number of Words by Message:', 'Average Number of Words by Message: {}'.format(average_message_words))
        p.text = p.text.replace('Average Number of Letters by Message:', 'Average Number of Letters by Message: {}'.format(average_message_letters))
        p.text = p.text.replace('Average Number of Messages per Day:', 'Average Number of Messages per Day: {}'.format(average_message_day))
    
    doc.add_page_break()
    doc.add_paragraph('WhatsApp Analysis')
    for filename in os.listdir(folder):
        if filename.endswith(".png"):
            doc.add_picture(folder + "\\" + filename, width=Cm(17.94))    
    doc.save(folder + '\\WhatsApp Chat Report {}.docx'.format(datetime.now().strftime('%d_%m_%Y')))


# In[9]:


#Defining a function which gives back all the necessary information
#Function combines all the above defined functions

def whatsapp_chat_analyzer():
    chat = whatsapp_chat_processor() # Use the fuction to import
    basic_statistics = input('Do you want to see basic Statistics about your Chat? (Yes or No)') #instruction for user
    total_messages = chat.shape[0]
    media_messages = chat[chat['message'] == '<Multimedia omitido>'].shape[0]
    average_message_words = round(chat['Word_Count'].mean(),1)
    average_message_letters = round(chat['Letter_Count'].mean(),1)
    average_message_day = round(chat.groupby('date')['message'].count().mean(),1)
    
    #if user indicates "yes" display the following statistics
    if basic_statistics.lower() == 'yes': 
        print('\033[1mBasic Statistics of {} \033[0m \nTotal Messages: {} \nMedia Message: {} \nAverage Number of Words by Messages: {} \nAverage Number of Letters by Messages: {} \nAverage Number of Messages per Day: {}'.format(
            filepath,
            total_messages,
            media_messages,
            average_message_words,
            average_message_letters,
            average_message_day)
             )  
    else: #if user indicates no then it skips the basic statistics
        pass
    
    #Create a Folder for Plots and Analysis
    folder = filepath[:-len(filepath.split('\\')[-1])] + 'Whatsapp Analysis {}'.format(datetime.now().strftime('%d_%m_%Y')
                                                                                      )
    os.makedirs(folder, exist_ok=True)
    
    #Creat pictures with statistics and save it in the above created folcder
    #Total Chats per Day
    chats_per_day = chat.groupby('dateTime')['message'].count().plot(
        kind='line', figsize=(20,10), color='#A26360',title='Total Chats per Day').get_figure()
    chats_per_day.savefig(folder + '\\chats_per_day_{}.png'.format(datetime.now().strftime('%d_%m_%Y'))
                         )

    plt.close(chats_per_day)
    
    #Most active Weekdays by Messages
    weekday_grouped_msg =  (chat.set_index('weekday')['message']
                          .groupby(level=0)
                          .value_counts()
                          .groupby(level=0)
                          .sum()
                          .reset_index(name='count')
                           )

    most_active_weekday = px.line_polar(weekday_grouped_msg, r='count', theta='weekday', line_close=True
                                       )
    most_active_weekday.update_traces(fill='toself'
                                     )
    most_active_weekday.update_layout(
        polar=dict(
            radialaxis=dict(
                visible=True,)),
        showlegend=False,
        title_text = 'Most active Weekdays by Messages'
    )
    most_active_weekday.write_image(folder + '\\most_active_weekday_{}.png'
                                    .format(datetime.now()
                                            .strftime('%d_%m_%Y')), format='.png'
                                   )
    
    #Most active Day Time by Messages
    hour_grouped_msg =  (chat.set_index('hour')['message']
                              .groupby(level=0)
                              .value_counts()
                              .groupby(level=0)
                              .sum()
                              .reset_index(name='count')
                        )
    most_active_day_time = px.bar(hour_grouped_msg, x='hour', y='count',
                     labels={'hour':'24 Hour Period'}, 
                     height=400
                                 )
    most_active_day_time.update_traces(marker_color='#EDCC8B', marker_line_color='#D4A29C',
                      marker_line_width=1.5, opacity=0.6
                                      )
    most_active_day_time.update_layout(title_text='Total Messages by Hour of the Day')
    most_active_day_time.write_image(folder + '\\most_active_day_time_{}.png'
                                    .format(datetime.now()
                                            .strftime('%d_%m_%Y')), format='.png'
                                    )
    
    #Most Active days per Month by Messages
    grouped_by_month_and_day = chat.groupby(['month_sent', 'weekday'])['message'].value_counts().reset_index(name='count')

    pt = grouped_by_month_and_day.pivot_table(index= 'month_sent', columns= 'weekday', values='count').reindex(index=months, columns= days)
    active_days_per_month = px.imshow(pt,
                    labels=dict(x='Day of Week', y='Months', color='Count'),
                    x=days,
                    y=months
                   )
    active_days_per_month.update_layout(
        width = 700,
        height = 700,
        title_text='Most Active days per Month by Messages'
    )
    active_days_per_month.write_image(folder + '\\active_days_per_month_{}.png'
                                    .format(datetime.now()
                                            .strftime('%d_%m_%Y')), format='.png'
                                     )
    
    #Most Active Users by Messages
    authors = chat['username'].value_counts().plot(kind='barh',figsize=(20,10), color=['#D4A29C', '#E8B298', '#EDCC8B', '#BDD1C5', '#9DAAA2']).get_figure()
    authors.savefig(folder + '\\most_active_users_{}.png'.format(datetime.now().strftime('%d_%m_%Y')))
    plt.close(authors)
    
    #Most Common Words function (defined above) is executed
    #(1) Create a graph with the most common words (last picture in folder)
    #(2) Statistics per user (after asking user)
    most_common_words(chat=chat, folder=folder)
    
    #Word Cloud
    chat_word_cloud = chat[['message']].copy()
    chat_word_cloud['message']= chat_word_cloud['message'].apply(emoji_free_text)# apply fuction which is defined above to clean text from emojis
    chat_word_cloud['message']= chat_word_cloud['message'].apply(remove_urls)# apply fuction which is defined above to clean text from urls
    text = " ".join(review for review in chat_word_cloud.message.dropna())
    wordcloud = WordCloud(width = 3000, height = 2000, random_state=1, 
                          background_color='black', colormap='Set2', collocations=False,
                          stopwords = STOPWORDS).generate(text)
    plt.figure(figsize=(40,30))
    plt.imshow(wordcloud)
    plt.axis("off")
    plt.savefig(folder + '\\wordcloud_{}.png'.format(datetime.now().strftime('%d_%m_%Y')))
    plt.close()
    
     #Create a Word report with the statistics and save it in the created folder - based on fuction defined above
    create_report_docx(total_messages, media_messages, average_message_words, average_message_letters, average_message_day, filepath, folder)
    


# In[10]:


whatsapp_chat_analyzer()


# In[ ]:




